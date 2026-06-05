#!/usr/bin/env python3
"""Fill DFSS Pilot Measurement Word template — P05 SecuGuard from HQ FILLIN markdown data."""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[2]
HQ_DFSS = REPO / "hq" / "07_DFSS"
TEMPLATE = Path(
    "/Users/marwanmahra/Library/CloudStorage/"
    "OneDrive-Cert-Expert/QM/Strategie/"
    "DFSS_PILOT_MEASUREMENT_ACTIVATION_TEMPLATE_V1_Cert_Expert_BILINGUAL.docx"
)
OUT_ONEDRIVE = TEMPLATE.parent / (
    "DFSS_PILOT_MEASUREMENT_ACTIVATION_P05_SecuGuard_FILLED_2026-06-05.docx"
)
OUT_HQ = HQ_DFSS / "DFSS_PILOT_MEASUREMENT_ACTIVATION_P05_SecuGuard_FILLED.docx"

STAND = date.today().isoformat()


def set_cell(cell, text: str) -> None:
    cell.text = text


def fill_header(doc: Document) -> None:
    t = doc.tables[0]
    set_cell(t.rows[7].cells[1], "TBD — Cert-Expert (sg03 Koordination)")
    set_cell(t.rows[8].cells[1], STAND)


def fill_quick_status(doc: Document) -> None:
    t = doc.tables[1]
    set_cell(
        t.rows[1].cells[1],
        "[ ] Yes  [x] No  [ ] Partly — NC counts yes; V-01–V-03 hours TBD",
    )
    set_cell(
        t.rows[2].cells[1],
        "[ ] Yes  [ ] No  [x] Partly — measurable since 2026-05-28 finding",
    )
    set_cell(
        t.rows[3].cells[1],
        "[ ] Yes  [x] No, Data Collection Required",
    )
    set_cell(
        t.rows[4].cells[1],
        "[x] Activate measurement system  [ ] Rework  [ ] Report later",
    )


def fill_portfolio_p05(doc: Document) -> None:
    row = doc.tables[2].rows[5]
    vals = [
        "P05",
        "Surveillance audit / NC closure (Ü1)",
        "ISO 45001, ISO 14001",
        "SecuGuard GmbH, Berlin (1 site, post-relocation)",
        "Continuity / Ü1 follow-up with customer",
        "2026-05-28",
        "2026-05-28 (finding); closure deadline 2026-06-30",
        "7 NCs open; Ü1 incomplete; EV-SG-001",
        "[ ] Green  [ ] Yellow  [x] Red",
    ]
    for i, v in enumerate(vals):
        set_cell(row.cells[i], v)


def fill_setup_record(doc: Document) -> None:
    t = doc.tables[3]
    data = {
        1: "P05",
        2: "P05-NC-closure-2026 (internal: SecuGuard)",
        3: "[ ] Initial cert  [x] Surveillance  [x] Other: NC closure after Ü1",
        4: "[ ] ISO 9001  [ ] DIN 77200-1/2  [x] ISO 14001  [x] ISO 45001",
        5: "Company-wide Berlin; admin site after relocation (env. m2)",
        6: "n/a",
        7: "[ ] Documentation basis  [ ] Done-for-you  [ ] Done-with-you  [x] Continuity",
        8: "2026-05-28",
        9: "2026-05-28 (finding); closure 2026-06-30",
        10: "5 major + 2 minor NCs from audit 28.05.2026; 0 closed",
        11: "HQ + OneDrive Ü1; Dashboard ToDos sg01–sg08",
        12: "TBD (Cert-Expert); customer: M. Marquardt",
    }
    for idx, text in data.items():
        set_cell(t.rows[idx].cells[1], text)


def fill_timeline_p05(doc: Document) -> None:
    t = doc.tables[4]
    col = 5
    comments = {
        1: "TBD",
        2: "Measurement phase / NC start",
        3: "TBD",
        4: "Ü1 deviation doc from audit",
        5: "TBD",
        6: "No — 0/7 NC closed",
        7: "2026-06-03 HQ import Audit_2026.md",
        8: "Partial — Ü1 in review",
        9: "n/a — closure phase",
        10: "2026-05-28",
        11: "7 NCs open (5+2)",
    }
    values = {
        1: "TBD",
        2: "2026-05-28",
        3: "TBD",
        4: "2026-05-28",
        5: "TBD",
        6: "No",
        7: "2026-06-03",
        8: "Partial",
        9: "n/a",
        10: "2026-05-28",
        11: "7 NCs open",
    }
    for r, val in values.items():
        set_cell(t.rows[r].cells[col], val)
        if r in comments:
            set_cell(t.rows[r].cells[6], comments[r])


def fill_blockers(doc: Document) -> None:
    t = doc.tables[5]
    rows = [
        (
            "B-01",
            "P05",
            "M1 Mutterschutz instruction missing",
            "S0-03, S0-05",
            "M. Marquardt",
            "2026-06-30",
            "[x] Open",
            "EV-SG-001",
        ),
        (
            "B-02",
            "P05",
            "M2 worker participation — verification pending",
            "S0-03, S0-05",
            "M. Marquardt",
            "2026-06-30",
            "[x] Open",
            "EV-SG-001",
        ),
        (
            "B-03",
            "P05",
            "M3 FASI/physician site-specific + annual reports",
            "S0-03",
            "TBD",
            "2026-06-30",
            "[x] Open",
            "EV-SG-001",
        ),
        (
            "B-04",
            "P05",
            "M4 supervisor monitoring evidence",
            "S0-03",
            "TBD",
            "2026-06-30",
            "[x] Open",
            "EV-SG-001",
        ),
        (
            "B-05",
            "P05",
            "M5 contract for sampled project",
            "S0-03",
            "TBD",
            "2026-06-30",
            "[x] Open",
            "EV-SG-001",
        ),
        (
            "B-06",
            "P05",
            "m1 prokuristin HR records",
            "S0-03",
            "TBD",
            "2026-06-30",
            "[x] Open",
            "EV-SG-002",
        ),
        (
            "B-07",
            "P05",
            "m2 environmental baselines after move",
            "S0-03",
            "M. Marquardt",
            "2026-06-30",
            "[x] Open",
            "EV-SG-002",
        ),
        (
            "B-08",
            "P05",
            "Ü1 forms per NC incomplete",
            "S0-03, S0-08",
            "TBD / customer",
            "2026-06-25",
            "[x] Open",
            "EV-SG-003",
        ),
    ]
    for i, vals in enumerate(rows, start=1):
        for j, v in enumerate(vals):
            set_cell(t.rows[i].cells[j], v)


def fill_evidence(doc: Document) -> None:
    t = doc.tables[6]
    rows = [
        (
            "EV-SG-001",
            "P05",
            "Major NCs M1–M5",
            "OneDrive Ü1 docx",
            "TBD",
            "TBD",
            "Document review",
            "in review",
            "No",
        ),
        (
            "EV-SG-002",
            "P05",
            "Minor NCs m1–m2",
            "OneDrive Ü1",
            "TBD",
            "TBD",
            "Document review",
            "in review",
            "No",
        ),
        (
            "EV-SG-003",
            "P05",
            "Ü1 forms per NC",
            "OneDrive Ü1",
            "TBD",
            "M. Marquardt",
            "Form completion",
            "in review",
            "No",
        ),
    ]
    for i, vals in enumerate(rows, start=1):
        for j, v in enumerate(vals):
            if j < len(t.rows[i].cells):
                set_cell(t.rows[i].cells[j], v)


def fill_baseline_p05(doc: Document) -> None:
    t = doc.tables[7]
    col = 6
    ev_col = 7
    data = [
        ("TBD", "EV-"),
        ("TBD", "EV-"),
        ("7 NCs + incomplete Ü1", "EV-SG-001/002/003"),
        ("Rework (NC closure)", "EV-SG-001"),
        ("7 open points", "EV-SG-001"),
        ("TBD", "EV-"),
        ("Deadline 2026-06-30; 8 days since audit", "EV-SG-003"),
    ]
    for i, (val, ev) in enumerate(data, start=1):
        set_cell(t.rows[i].cells[col], val)
        set_cell(t.rows[i].cells[ev_col], ev)


def fill_s0(doc: Document) -> None:
    t = doc.tables[8]
    statuses = [
        ("[x] Fail", "EV-SG-001", "7 NCs NOT READY"),
        ("[x] Fail", "—", "SDL detail TBD"),
        ("[x] Fail", "EV-SG-001/002/003", "0 verified closure"),
        ("[ ] TBD", "—", "Ü1 not released"),
        ("[x] Fail", "EV-SG-001", "many owners TBD"),
        ("[ ] TBD", "—", "FASI/physician M3"),
        ("[x] Fail", "EV-SG-001", "evidence missing"),
        ("[x] Fail", "EV-SG-003", "Ü1 forms open"),
    ]
    for i, (st, ev, _note) in enumerate(statuses, start=1):
        set_cell(t.rows[i].cells[4], st)
        set_cell(t.rows[i].cells[5], ev)


def fill_readiness_p05(doc: Document) -> None:
    row = doc.tables[9].rows[5]
    set_cell(row.cells[1], "[ ]")
    set_cell(row.cells[2], "[ ]")
    set_cell(row.cells[3], "[x]")
    set_cell(
        row.cells[4],
        "7 NCs open; 0% closure; S0 fail",
    )
    set_cell(
        row.cells[5],
        "EV-SG-001 — action plan M1–M5 per NC",
    )


def fill_pilot_v(doc: Document) -> None:
    t = doc.tables[10]
    rows = [
        ("P05", "TBD", "TBD", "EV-", "[x] Open"),
        ("P05", "n/a closure phase", "—", "—", "TBD"),
        ("P05", "TBD", "TBD", "EV-", "[x] Open"),
        (
            "P05",
            "5 Major + 2 Minor open",
            "TBD",
            "EV-SG-001",
            "[x] Captured",
        ),
        ("P05", "TBD", "TBD", "EV-", "[x] Open"),
        ("P05", "TBD", "TBD", "EV-", "[x] Open"),
        (
            "P05",
            "Ü1 imported; verification pending",
            "TBD",
            "EV-SG-003",
            "[ ] Partly",
        ),
        ("P05", "Incomplete", "TBD", "EV-SG-001", "[x] Captured"),
        (
            "P05",
            "Deadline 2026-06-30; 8 days since audit",
            "TBD",
            "EV-SG-003",
            "[x] Captured",
        ),
        ("P05", "TBD (relocation, cross-cutting)", "TBD", "EV-", "[x] Open"),
    ]
    for i, vals in enumerate(rows, start=1):
        set_cell(t.rows[i].cells[2], vals[0])
        set_cell(t.rows[i].cells[3], vals[1])
        set_cell(t.rows[i].cells[4], vals[2])
        set_cell(t.rows[i].cells[5], vals[3])
        set_cell(t.rows[i].cells[6], vals[4])


def fill_gate(doc: Document) -> None:
    t = doc.tables[16]
    decisions = [
        ("[x] No", "EV-SG-001", "0/7 NC closed"),
        ("[x] No", "§7 S0", "majority Fail"),
        ("[x] No", "—", "RC not started"),
        ("[ ] Partly", "§13", "structure yes, V-values no"),
        ("[x] DATA COLLECTION REQUIRED", "—", "measurement active"),
    ]
    for i, (dec, ev, com) in enumerate(decisions, start=1):
        set_cell(t.rows[i].cells[1], dec)
        set_cell(t.rows[i].cells[2], ev)
        set_cell(t.rows[i].cells[3], com)


def fill_data_quality(doc: Document) -> None:
    t = doc.tables[14]
    if len(t.rows) < 4:
        return
    answers = [
        "HQ customer file + OneDrive Ü1 (not full text in git)",
        "No — documented facts only",
        "No hours V-01–V-03; internal owner TBD",
        "Capture hours at sg03 meetings; first NC close → DFSS usable Yes",
    ]
    for i, a in enumerate(answers, start=1):
        if i < len(t.rows):
            set_cell(t.rows[i].cells[1], a)


def main() -> None:
    if not TEMPLATE.is_file():
        raise SystemExit(f"Template not found: {TEMPLATE}")

    shutil.copy2(TEMPLATE, OUT_ONEDRIVE)
    doc = Document(str(OUT_ONEDRIVE))

    fill_header(doc)
    fill_quick_status(doc)
    fill_portfolio_p05(doc)
    fill_setup_record(doc)
    fill_timeline_p05(doc)
    fill_blockers(doc)
    fill_evidence(doc)
    fill_baseline_p05(doc)
    fill_s0(doc)
    fill_readiness_p05(doc)
    fill_pilot_v(doc)
    fill_data_quality(doc)
    fill_gate(doc)

    doc.save(str(OUT_ONEDRIVE))
    shutil.copy2(OUT_ONEDRIVE, OUT_HQ)

    print(f"Saved: {OUT_ONEDRIVE}")
    print(f"Copy:  {OUT_HQ}")


if __name__ == "__main__":
    main()
