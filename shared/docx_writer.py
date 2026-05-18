from docx import Document
from docx.shared import Pt

PLACEHOLDER_LABELS = {
    "GB_TAETIGKEIT": "Tätigkeit / Einsatzbeschreibung",
    "GB_GEFAEHRDUNGEN": "Identifizierte Gefährdungen",
    "GB_RISIKOBEWERTUNG": "Risikobewertung",
    "GB_SCHUTZMASSNAHMEN": "Schutzmaßnahmen",
    "GB_OFFENE_PUNKTE": "Offene Punkte (Zusammenfassung)",
}


def save_docx(title: str, content: str, output_path: str):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(output_path)
    print(f"DOCX gespeichert: {output_path}")


def save_structured_docx(output: dict, output_path: str):
    """
    Write a structured bot output dict to a formatted DOCX test document.

    Prefixes the document title with [DRAFT] if qa_status is not 'ok'.
    Each placeholder is rendered as a labeled section. Open points are
    appended as a bullet list at the end.

    Args:
        output:      Enriched output dict from quality_checker.check().
        output_path: Destination file path for the DOCX.
    """
    doc = Document()

    qa_status = output.get("qa_status", "review_required")
    doc_type = output.get("document_type", "Dokument")
    meta = output.get("meta", {})
    event_name = meta.get("event_name", "")

    if qa_status != "ok":
        title = f"[DRAFT] Gefährdungsbeurteilung — {event_name}".strip(" —")
    else:
        title = f"Gefährdungsbeurteilung — {event_name}".strip(" —")

    doc.add_heading(title, level=1)

    status_para = doc.add_paragraph()
    status_run = status_para.add_run(f"QA-Status: {qa_status.upper()}")
    status_run.bold = True
    if qa_status != "ok":
        status_run.font.size = Pt(11)

    doc.add_paragraph(
        f"Erstellt: {meta.get('created_at', '')}   |   "
        f"Version: {meta.get('pipeline_version', '')}   |   "
        f"Ersteller: {meta.get('created_by', '')}"
    )

    doc.add_paragraph("")

    placeholders = output.get("placeholders", {})
    for key, value in placeholders.items():
        label = PLACEHOLDER_LABELS.get(key, key)
        doc.add_heading(label, level=2)
        doc.add_paragraph(value if value else "—")

    open_points = output.get("open_points", [])
    if open_points:
        doc.add_heading("Offene Punkte", level=2)
        for point in open_points:
            doc.add_paragraph(point, style="List Bullet")

    doc.save(output_path)
    print(f"DOCX gespeichert: {output_path}")