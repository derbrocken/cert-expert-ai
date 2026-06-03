"""EC template generator — Pflichtform for ec_event_kampfsport."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = Path("templates/ec_event_kampfsport.docx")


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.clear()
    header_run = header.add_run("{CompanyName}  |  ")
    header_run.bold = True
    header.add_run("{Logo}")

    doc.add_heading("Einsatzkonzept — Kampfsportveranstaltung", level=0)

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    for i, (label, val) in enumerate(
        [
            ("Dokumentversion", "{DocVersion}"),
            ("Erstellt am", "{DocDate}"),
            ("Erstellt von", "{CreatedBy}"),
            ("Freigegeben von", "{ApprovedBy}"),
            ("Auftraggeber / Veranstalter", "{ClientName}"),
        ]
    ):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = val

    doc.add_paragraph(
        "Operatives Einsatzkonzept aus Sicht des Sicherheitsdienstleisters (AN). "
        "Kein Ersatz für Sicherheitskonzept (SK) oder Gefährdungsbeurteilung (GB)."
    )

    for title, key in [
        ("1. Einsatzbeschreibung", "{EC_EINSATZBESCHREIBUNG}"),
        ("2. Kräfte und Rollen", "{EC_KRAEFTE_UND_ROLLEN}"),
        ("3. Positionierung und Abschnitte", "{EC_POSITIONIERUNG_UND_ABSCHNITTE}"),
        ("4. Ablauf und Briefing", "{EC_ABLAUF_UND_BRIEFING}"),
        ("5. Kommunikation und Meldewege", "{EC_KOMMUNIKATION_UND_MELDEWEGE}"),
        ("6. Logistik / Ausrüstung", "{EC_LOGISTIK_AUSRUESTUNG}"),
        ("7. Notfallplan", "{EC_NOTFALLPLAN}"),
        ("8. Dokumentation / Nachbereitung", "{EC_DOKUMENTATION_NACHBEREITUNG}"),
        ("9. Offene Punkte", "{EC_OFFENE_PUNKTE}"),
    ]:
        doc.add_heading(title, level=1)
        doc.add_paragraph(key)

    doc.add_heading("Unterschriften", level=1)
    doc.add_paragraph("Erstellt von: {CreatedBy}  |  Datum: {DocDate}")
    doc.add_paragraph("Freigegeben von: {ApprovedBy}  |  Datum: __________")

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.add_run("Cert-Expert AI — Einsatzkonzept  |  {DocVersion}  |  {DocDate}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(str(OUTPUT_PATH))
    print(f"Template erstellt: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

