"""
Template generator for the GB blueprint `gb_event_kampfsport`.

Produces `templates/gb_event_kampfsport.docx` with:
  - Layer 1 placeholders ({CompanyName}, {Logo}, {DocVersion}, {DocDate},
    {CreatedBy}, {ApprovedBy}, {currentDate})
  - Layer 3 GB placeholders for this blueprint:
      {GB_TAETIGKEIT}
      {GB_GEFAEHRDUNGEN}
      {GB_RISIKOBEWERTUNG}
      {GB_SCHUTZMASSNAHMEN}
      {GB_VERANTWORTLICHKEITEN}
      {GB_OFFENE_PUNKTE}
  - Chapter structure aligned with `docs/BLUEPRINT_ARCHITECTURE.md` §4.2

The template is the static side of the document — chapter headings,
the rating-matrix legend, the STOP-Prinzip note, and the signature
section are owned by Cert-Expert (template authors), not by the bot.

Run from project root:
    python3 templates/create_gb_event_kampfsport_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = Path("templates/gb_event_kampfsport.docx")


def add_labeled(doc: Document, label: str, token: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(token)


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.clear()
    run = header_para.add_run("{CompanyName}  |  ")
    run.bold = True
    header_para.add_run("{Logo}")
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("Gefährdungsbeurteilung — Kampfsportveranstaltung", level=0)

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_rows = [
        ("Dokumentversion",   "{DocVersion}"),
        ("Erstellt am",       "{DocDate}"),
        ("Erstellt von",      "{CreatedBy}"),
        ("Freigegeben von",   "{ApprovedBy}"),
    ]
    for i, (label, value) in enumerate(meta_rows):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value

    doc.add_paragraph("")
    doc.add_paragraph(
        "Diese Gefährdungsbeurteilung wird im Rahmen des Arbeitsschutzgesetzes "
        "und der Unfallverhütungsvorschriften erstellt. Sie ist ein "
        "Arbeitsdokument und wird durch die im Feld \u201eFreigegeben von\u201c "
        "benannte verantwortliche Person geprüft und freigegeben."
    )

    doc.add_heading("1. Einsatzbeschreibung", level=1)
    doc.add_paragraph("{GB_TAETIGKEIT}")

    doc.add_heading("2. Veranstaltungsspezifische Gefährdungsanalyse", level=1)
    doc.add_paragraph(
        "Im Folgenden werden die identifizierten Gefährdungen für die unter "
        "Abschnitt 1 beschriebene Veranstaltung dargestellt. Die Gliederung "
        "folgt den Kategorien Publikum, Kampfsport, Örtlichkeit und Personal."
    )
    doc.add_paragraph("{GB_GEFAEHRDUNGEN}")

    doc.add_heading("3. Risikobewertung", level=1)
    doc.add_paragraph(
        "Bewertet wird nach Eintrittwahrscheinlichkeit (gering / mittel / hoch) "
        "und Schadensschwere (gering / mittel / hoch). Daraus ergibt sich ein "
        "Gesamtrisiko (niedrig / mittel / hoch). Numerische Risikokennzahlen "
        "werden nicht angegeben."
    )
    doc.add_paragraph("{GB_RISIKOBEWERTUNG}")

    doc.add_heading("4. Schutzmaßnahmen", level=1)
    doc.add_paragraph(
        "Maßnahmen werden nach dem STOP-Prinzip strukturiert: "
        "Substitution, Technisch, Organisatorisch, Persönlich. "
        "Persönliche Schutzmaßnahmen folgen nachrangig auf technische und "
        "organisatorische Maßnahmen."
    )
    doc.add_paragraph("{GB_SCHUTZMASSNAHMEN}")

    doc.add_heading("5. Verantwortlichkeiten", level=1)
    doc.add_paragraph("{GB_VERANTWORTLICHKEITEN}")

    doc.add_heading("6. Offene Punkte", level=1)
    doc.add_paragraph(
        "Folgende Punkte sind vor der Freigabe dieses Dokuments zu klären. "
        "Solange offene Punkte bestehen, gilt das Dokument als Entwurf."
    )
    doc.add_paragraph("{GB_OFFENE_PUNKTE}")

    doc.add_heading("Unterschriften", level=1)
    doc.add_paragraph("Erstellt von: {CreatedBy}")
    doc.add_paragraph("Datum: {DocDate}")
    doc.add_paragraph("")
    doc.add_paragraph("Geprüft und freigegeben von: {ApprovedBy}")
    doc.add_paragraph("Datum: ____________________________")

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.add_run(
        "Cert-Expert AI — Gefährdungsbeurteilung Kampfsport  |  "
        "Version: {DocVersion}  |  Datum: {DocDate}  |  "
        "Erstellt: {CreatedBy}  |  Freigabe: {ApprovedBy}  |  "
        "Auto-Datum: {currentDate}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    doc = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Template erstellt: {OUTPUT_PATH}")
    print()
    print("Enthaltene Platzhalter:")
    print("  Layer 1 (Header): {CompanyName}, {Logo}")
    print("  Layer 1 (Meta):   {DocVersion}, {DocDate}, {CreatedBy}, {ApprovedBy}")
    print("  Layer 1 (Footer): {DocVersion}, {DocDate}, {CreatedBy},"
          " {ApprovedBy}, {currentDate}")
    print("  Layer 3 (GB):     {GB_TAETIGKEIT}, {GB_GEFAEHRDUNGEN},")
    print("                    {GB_RISIKOBEWERTUNG}, {GB_SCHUTZMASSNAHMEN},")
    print("                    {GB_VERANTWORTLICHKEITEN}, {GB_OFFENE_PUNKTE}")


if __name__ == "__main__":
    main()
