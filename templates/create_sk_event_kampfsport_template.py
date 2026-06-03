"""SK template generator — Pflichtform for sk_event_kampfsport."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = Path("templates/sk_event_kampfsport.docx")


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    run = header.add_run("{CompanyName}  |  ")
    run.bold = True
    header.add_run("{Logo}")

    doc.add_heading("Sicherheitskonzept — Kampfsportveranstaltung", level=0)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Dokumentversion", "{DocVersion}"),
        ("Erstellt am", "{DocDate}"),
        ("Erstellt von", "{CreatedBy}"),
        ("Freigegeben von", "{ApprovedBy}"),
    ]):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = val

    doc.add_paragraph(
        "Organisatorisches Sicherheitskonzept aus Veranstalter-/Gesamtperspektive. "
        "Kein Ersatz für die arbeitsschutzliche Gefährdungsbeurteilung (GB) "
        "oder das Einsatzkonzept (EK)."
    )

    doc.add_heading("1. Schutzziel", level=1)
    doc.add_paragraph("{SK_SCHUTZZIEL}")

    doc.add_heading("2. Gefährdungsanalyse", level=1)
    doc.add_paragraph("{SK_GEFAEHRDUNGSANALYSE}")

    doc.add_heading("3. Schutzmaßnahmen", level=1)
    doc.add_paragraph("{SK_SCHUTZMASSNAHMEN}")

    doc.add_heading("4. Verantwortliche", level=1)
    doc.add_paragraph("{SK_VERANTWORTLICHE}")

    doc.add_heading("5. Kommunikation und Meldewesen", level=1)
    doc.add_paragraph("{SK_KOMMUNIKATION}")

    doc.add_heading("6. Offene Punkte", level=1)
    doc.add_paragraph("{SK_OFFENE_PUNKTE}")

    doc.add_heading("Unterschriften", level=1)
    doc.add_paragraph("Erstellt von: {CreatedBy}  |  Datum: {DocDate}")
    doc.add_paragraph("Freigegeben von: {ApprovedBy}  |  Datum: __________")

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.add_run(
        "Cert-Expert AI — Sicherheitskonzept Kampfsport  |  {DocVersion}  |  {DocDate}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(str(OUTPUT_PATH))
    print(f"Template erstellt: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
