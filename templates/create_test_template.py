"""
Regression test template generator for the Cert-Expert AI renderer.

Creates templates/test_render.docx with placeholders from all three layers:
  - Layer 1: Company / document metadata (Tool 1 convention)
  - Layer 3: Fachbot content (GB namespace)
  - Edge cases: {currentDate} lowercase, empty {ApprovedBy}, {Logo}

Covers: body paragraphs, table cells, header, footer.

Run from project root:
    python3 templates/create_test_template.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


OUTPUT_PATH = Path("templates/test_render.docx")


def add_heading(doc: Document, text: str, level: int = 2):
    p = doc.add_heading(text, level=level)
    return p


def add_labeled_paragraph(doc: Document, label: str, token: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(token)
    return p


def build_document() -> Document:
    doc = Document()

    # ── Header ──────────────────────────────────────────────────────────────
    # Contains: {CompanyName} and {Logo} as plain text placeholders.
    # easy-template-x will replace them in-place (TextNode mode).
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    run = header_para.add_run("{CompanyName}  |  ")
    run.bold = True
    header_para.add_run("{Logo}")
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("RENDER-TEST — Cert-Expert AI Renderer", level=1)
    doc.add_paragraph(
        "Dieses Dokument ist das Regressions-Template für den Cert-Expert AI Renderer. "
        "Es enthält Platzhalter aus allen Ebenen: Unternehmens-Metadaten, "
        "Fachbot-Inhalte (GB-Namespace) und Randfälle."
    )

    # ── Section 1: Company / Document Metadata (Layer 1) ─────────────────────
    add_heading(doc, "1. Unternehmens- und Dokumentdaten (Layer 1)", level=2)
    add_labeled_paragraph(doc, "Unternehmen", "{CompanyName}")
    add_labeled_paragraph(doc, "Straße", "{CompanyStreet}")
    add_labeled_paragraph(doc, "PLZ / Stadt", "{CompanyZip} {CompanyCity}")
    add_labeled_paragraph(doc, "Land", "{CompanyCountry}")
    add_labeled_paragraph(doc, "Adresszusatz", "{CompanyAddressLine}")
    add_labeled_paragraph(doc, "E-Mail", "{CompanyEmail}")
    add_labeled_paragraph(doc, "Erstellt am (auto)", "{currentDate}")

    # ── Section 2: Document Metadata in a Table ───────────────────────────────
    add_heading(doc, "2. Dokumentmetadaten (Tabelle)", level=2)
    doc.add_paragraph(
        "Diese Tabelle prüft Platzhalterersetzung in Tabellenzellen."
    )

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    rows_data = [
        ("Feld", "Wert"),
        ("Erstellt von", "{CreatedBy}"),
        ("Freigegeben von", "{ApprovedBy}"),
        ("Dokumentversion", "{DocVersion}"),
        ("Dokumentdatum", "{DocDate}"),
    ]

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        cell_label.text = label
        cell_value.text = value
        if i == 0:
            for cell in (cell_label, cell_value):
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph("")

    # ── Section 3: Logo test ──────────────────────────────────────────────────
    add_heading(doc, "3. Logo-Ersetzung (Layer 1)", level=2)
    doc.add_paragraph(
        "Der folgende Platzhalter {Logo} wird durch das Unternehmenslogo ersetzt. "
        "Falls kein Logo angegeben, erscheint an dieser Stelle nichts."
    )
    p = doc.add_paragraph()
    p.add_run("Logo: ")
    p.add_run("{Logo}")

    # ── Section 4: GB Bot Content (Layer 3) ───────────────────────────────────
    add_heading(doc, "4. Gefährdungsbeurteilung — Fachbot-Inhalte (Layer 3)", level=2)
    doc.add_paragraph(
        "Dieser Abschnitt enthält die GB-Namespace-Platzhalter, "
        "die vom Gefährdungsbeurteilungs-Bot befüllt werden."
    )
    add_labeled_paragraph(doc, "Tätigkeit", "{GB_TAETIGKEIT}")
    add_labeled_paragraph(doc, "Gefährdungen", "{GB_GEFAEHRDUNGEN}")
    add_labeled_paragraph(doc, "Risikobewertung", "{GB_RISIKOBEWERTUNG}")
    add_labeled_paragraph(doc, "Schutzmaßnahmen", "{GB_SCHUTZMASSNAHMEN}")
    add_labeled_paragraph(doc, "Offene Punkte", "{GB_OFFENE_PUNKTE}")

    # ── Section 5: Edge Cases ─────────────────────────────────────────────────
    add_heading(doc, "5. Randfälle", level=2)

    doc.add_paragraph(
        "5a. Leerstring-Ersetzung — {ApprovedBy} ist leer, "
        "wenn keine Freigabe eingetragen ist. "
        "An dieser Stelle darf kein sichtbares Token erscheinen:"
    )
    p = doc.add_paragraph()
    p.add_run("Freigabe: [").bold = False
    p.add_run("{ApprovedBy}")
    p.add_run("]")

    doc.add_paragraph(
        "5b. Kleinbuchstaben-Randfall — {currentDate} hat ein kleines 'c'. "
        "Dieser Token muss trotzdem korrekt ersetzt werden:"
    )
    add_labeled_paragraph(doc, "Datum (auto)", "{currentDate}")

    doc.add_paragraph(
        "5c. Nicht befüllter Platzhalter — {GB_RISIKOBEWERTUNG} erscheint "
        "nochmals hier, um zu prüfen dass derselbe Token an mehreren Stellen "
        "konsistent ersetzt wird:"
    )
    add_labeled_paragraph(doc, "Risikobewertung (Wiederholung)", "{GB_RISIKOBEWERTUNG}")

    # ── Footer ────────────────────────────────────────────────────────────────
    # Contains: {DocVersion}, {DocDate}, {CreatedBy}, {ApprovedBy}
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.add_run("Version: {DocVersion}  |  Datum: {DocDate}  |  "
                        "Erstellt: {CreatedBy}  |  Freigabe: {ApprovedBy}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    doc = build_document()
    doc.save(str(OUTPUT_PATH))
    print(f"Regression-Template erstellt: {OUTPUT_PATH}")
    print()
    print("Enthaltene Platzhalter:")
    print("  Layer 1 (Kopfzeile): {CompanyName}, {Logo}")
    print("  Layer 1 (Fußzeile):  {DocVersion}, {DocDate}, {CreatedBy}, {ApprovedBy}")
    print("  Layer 1 (Body):      {CompanyName}, {CompanyStreet}, {CompanyZip},")
    print("                       {CompanyCity}, {CompanyCountry}, {CompanyAddressLine},")
    print("                       {CompanyEmail}, {currentDate}")
    print("  Layer 1 (Tabelle):   {CreatedBy}, {ApprovedBy}, {DocVersion}, {DocDate}")
    print("  Layer 1 (Logo):      {Logo}")
    print("  Layer 3 (GB):        {GB_TAETIGKEIT}, {GB_GEFAEHRDUNGEN},")
    print("                       {GB_RISIKOBEWERTUNG}, {GB_SCHUTZMASSNAHMEN},")
    print("                       {GB_OFFENE_PUNKTE}")
    print("  Edge cases:          {ApprovedBy} (leer), {currentDate} (klein),")
    print("                       {GB_RISIKOBEWERTUNG} (doppelt)")


if __name__ == "__main__":
    main()
